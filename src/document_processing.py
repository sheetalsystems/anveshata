# src/document_processing.py
import os
import pickle
import PyPDF2
from docx import Document
import time
import faiss
from sentence_transformers import SentenceTransformer
import numpy as np
from pdf2image import convert_from_path
import pytesseract
import concurrent.futures
import google.generativeai as genai
from google.generativeai import GenerativeModel, configure
from src.config import GEN_AI_API_KEY, GENERATIVE_MODEL, EMBEDDING_MODEL
from src.onedrive_integration import get_sharepoint_file_content, get_sharepoint_file_ids
import asyncio
import docx
import requests
import json
import io
import aiohttp
from urllib.parse import quote
import pandas as pd
import re

# Configure the Gemini API key.
configure(api_key=GEN_AI_API_KEY)  # Replace with your actual API key.
model = GenerativeModel(GENERATIVE_MODEL)
embedding_model = SentenceTransformer(EMBEDDING_MODEL)
cached_documents = {}

all_chunks = None
file_id_map = None

embedding_dimension = 768  
faiss_index = faiss.IndexHNSWFlat(embedding_dimension, 32)
faiss_index.hnsw.efConstruction = 200
faiss_index.hnsw.efSearch = 50

def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
                try:
                    images = convert_from_path(file_path, first_page=page.page_number + 1, last_page=page.page_number + 1)
                    if images:
                        for img in images:
                            text += pytesseract.image_to_string(img)
                except:
                    pass
            return text
    except Exception as e:
        return f"Error reading PDF file: {e}"

def read_word(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        return f"Error reading Word file: {e}"

def chunk_text(text, chunk_size=5000):
    words = text.split()
    return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]



async def load_documents(client_id, client_secret, tenant_id, site_id):
    """Loads .docx, .pdf, and .xlsx documents from SharePoint, extracts and chunks their content."""
    global cached_documents
    cached_documents = {}

    file_items = await get_sharepoint_file_ids(client_id, client_secret, tenant_id, site_id)

    if not file_items or not file_items.get("value"):
        print("No files found or error retrieving file list.")
        return None  # Explicitly return None to avoid unpack error

    for item in file_items["value"]:
        if "file" not in item:
            continue  # Skip folders or non-files

        mime_type = item["file"].get("mimeType", "")
        download_url = item["@microsoft.graph.downloadUrl"]
        file_name = item.get("name", "unknown")

        try:
            response = requests.get(download_url)
            response.raise_for_status()
            content = response.content

            full_text = []

            # DOCX
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(io.BytesIO(content))
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        if any(row_text):
                            full_text.append(" | ".join(row_text))
                print(f"Processed DOCX: {file_name}")

            # PDF
            elif mime_type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        full_text.append(page_text.strip())
                print(f"Processed PDF: {file_name}")

            # Excel
            elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                try:
                    excel_data = pd.ExcelFile(io.BytesIO(content))
                    for sheet_name in excel_data.sheet_names:
                        df = excel_data.parse(sheet_name)
                        for row in df.itertuples(index=False):
                            row_text = [str(cell).strip() for cell in row if pd.notnull(cell)]
                            if row_text:
                                full_text.append(" | ".join(row_text))
                    print(f"Processed Excel: {file_name}")
                except Exception as e:
                    print(f"Error processing Excel {file_name}: {e}")
                    continue

            else:
                print(f"Skipping unsupported file type: {mime_type} for {file_name}")
                continue

            text = "\n".join(full_text)
            chunk_size = 5000
            chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
            cached_documents[file_name] = chunks
            print(f"Loaded and chunked: {file_name} into {len(chunks)} chunks")

        except requests.exceptions.RequestException as e:
            print(f"Error downloading {file_name}: {e}")
        except Exception as e:
            print(f"Error processing {file_name}: {e}")

    # Build combined chunk list and file_id map
    all_chunks = []
    file_id_map = {}

    for doc_name, chunks in cached_documents.items():
        for chunk in chunks:
            file_id_map[len(all_chunks)] = doc_name
            all_chunks.append(chunk)

    return all_chunks, file_id_map


async def download_and_process_doc(session, download_url, item_name, cached_documents):
    """Downloads a DOCX file and extracts text content, ignoring images."""
    try:
        async with session.get(download_url, timeout=30) as response:
            response.raise_for_status()
            content = await response.read()
            if content:
                try:
                    doc = docx.Document(io.BytesIO(content))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    cached_documents[item_name] = text
                    print(f"Loaded text content from: {item_name}")
                except docx.opc.exceptions.PackageNotFoundError:
                    print(f"Error: Invalid DOCX file - {item_name}")
                except Exception as e:
                    print(f"Error processing DOCX content of {item_name}: {e}")
            else:
                print(f"Warning: Empty content received for {item_name}")
    except aiohttp.ClientError as e:
        print(f"Error downloading {item_name}: {e}")
    except asyncio.TimeoutError:
        print(f"Timeout downloading {item_name}")
    except Exception as e:
        print(f"An unexpected error occurred with {item_name}: {e}")

def get_embedding(text):
    try:
        embedding = embedding_model.encode(text)
        return embedding.tolist()
    except Exception as e:
        print(f"Embedding generation error: {e}")
        return None
    
    

FOLLOW_UP_PHRASES = [
    "tell me more", "explain more", "related to it", "explain in detail", "of it", "more about it", "more about",
    "provide more insights", "provide more details", "provide more information", "elaborate",
    "give detail information about", "give more information about", "give more details about", "give more insights about",
    "give more insights on", "give more details on", "give more information on", "give more insights regarding",
    "give more details regarding", "give more information regarding", "give more insights related to",
    "give more details related to", "give more information related to"
]

def is_follow_up_question(question):
    q = question.lower()
    return any(phrase in q for phrase in FOLLOW_UP_PHRASES)


INDEX_PATH = "./faiss_hnsw.index"
META_PATH = "./faiss_metadata.pkl"

async def load_or_build_faiss_hnsw_index(client_id, client_secret, tenant_id, site_id,force_rebuild=False):
    global faiss_index, all_chunks, file_id_map
    
    if not force_rebuild and os.path.exists(INDEX_PATH) and os.path.exists(META_PATH):
        print("Loading FAISS index and metadata from disk...", flush=True)
        faiss_index = faiss.read_index(INDEX_PATH)
        with open(META_PATH, "rb") as f:
            metadata = pickle.load(f)
            all_chunks = metadata["all_chunks"]
            file_id_map = metadata["file_id_map"]
        print(f"Loaded FAISS index with {len(all_chunks)} chunks from disk.", flush=True)
        return
    
    print("Force rebuild enabled. Rebuilding FAISS index......", flush=True)
    all_chunks, file_id_map = await load_documents(client_id, client_secret, tenant_id, site_id)
    embeddings = [get_embedding(chunk) for chunk in all_chunks]
    embeddings_np = np.array(embeddings).astype('float32')

    # Load documents and chunk text
    # all_chunks, file_id_map = await load_documents(client_id, client_secret, tenant_id, site_id)

    # # Create embeddings for all chunks (assuming get_embedding is sync or async)
    # import numpy as np
    # embeddings = []
    # for chunk in all_chunks:
    #     emb = get_embedding(chunk)
    #     embeddings.append(emb)

    # embeddings_np = np.array(embeddings).astype('float32')

    # Create FAISS index and add embeddings
    embedding_dimension = embeddings_np.shape[1]
    faiss_index = faiss.IndexHNSWFlat(embedding_dimension, 32)
    faiss_index.hnsw.efConstruction = 200
    faiss_index.hnsw.efSearch = 50
    faiss_index.add(embeddings_np)

    faiss.write_index(faiss_index, INDEX_PATH)
    with open(META_PATH, "wb") as f:
        pickle.dump({"all_chunks": all_chunks, "file_id_map": file_id_map}, f)

    print(f"Built and saved FAISS index with {len(all_chunks)} chunks.", flush=True)









async def query_agent_with_faiss_hnsw(question, client_id, client_secret, tenant_id, site_id, chat_history=None):
 

    print("Entering FAISS-based query_agent_with_faiss_hnsw with question:", question, flush=True)

    # Ensure FAISS index is initialized
    if faiss_index is None or all_chunks is None or file_id_map is None:
        print("FAISS index not found in memory. Loading/rebuilding...", flush=True)
        await load_or_build_faiss_hnsw_index(client_id, client_secret, tenant_id, site_id)

        # return "System is not ready. Please try again later.", 0, 0, 0, None, False, False
    if faiss_index is None:
        return "FAISS index is not initialized.", 0, 0, 0, None, False, False
    

    is_follow_up = is_follow_up_question(question)

    try:
        question_embedding = get_embedding(question)
        if question_embedding is None:
            return "Could not generate embedding for the question.", 0, 0, 0, None, False, False
    except Exception as e:
        return f"Error getting question embedding: {e}", 0, 0, 0, None, False, False

    # Search top K matches
    k = 5
    question_embedding_np = np.array([question_embedding], dtype=np.float32)
    distances, indices = faiss_index.search(question_embedding_np, k)

    context = []
    most_relevant_document = None

    for i, idx in enumerate(indices[0]):
        if idx < 0 or idx >= len(all_chunks):
            continue
        chunk = all_chunks[idx]
        file_id = file_id_map[idx]
        if i == 0:
            most_relevant_document = file_id
        context.append(f"Document: {file_id}\n{chunk}")

    if not context:
        context.append("No relevant documents found.")

    # Prompt setup
    prompt = (
        # "You are an intelligent assistant helping IT support agents troubleshoot user issues using internal technical documentation.\n"
        # "Always assume the user is an IT support team member, not an end-user.\n\n"
        # "Given the query and the retrieved context:\n"
        # "- If the query is unclear or too broad, ask one clarifying question.\n"
        # "- Present 2–3 relevant options to guide the user (e.g., possible causes, tools to check, affected systems).\n"
        # "- If the query is clear, provide step-by-step troubleshooting steps or relevant document excerpts.\n"
        # "- Do not say 'contact IT support'; instead, provide helpful technical direction.\n"
        # "- Keep responses concise, clear, and structured (bullets or numbered steps).\n\n"
        
        "You are an intelligent assistant helping recruiter to find suitable profiles from provided job descriptions.\n"
        #"Always assume the user is a recruiter (not an IT or HR expert).\n"
        #"\n"
        "Given the query and the retrieved context:\n"
        "- If the query is unclear or too broad, ask one clarifying question.\n"
        "- Present 2–3 relevant options to guide the user (e.g., possible topics, resources, next steps).\n"
        "- If the query is clear, provide step-by-step answers or relevant document excerpts.\n"
        "- Do not say 'contact HR' or 'contact IT support'; instead, provide helpful technical or policy direction.\n"
        "- Keep responses concise, clear, and structured (bullets or numbered steps).\n"
        "- Use a friendly, professional tone.\n"
        "\n"
        "You can answer questions about:\n"
        "- profiles most suitable for job description\n"
        "= priority should be mandatory skills\n"
        "- compare between profiles and give maching percentages in a single, professional row html table without markdown code blocks and immediaate start with table tag based on these rules\n"
        "1. COLUMNS: Rank | Candidate Name | Match % | Experience |Core Skills Found | Primary Reason for Match\n"
        "2. SORTING: Sort by experience and Match % from Highest to Lowest.\n"
        "3. NO PREAMBLE: Do not say 'Here is your table.'\n"
    )

    if chat_history:
        for entry in chat_history[-2:]:
            prompt += f"User: {entry['question']}\nBot: {entry['answer']}\n---\n"

    if is_follow_up:
        prompt += "\nThis is a follow-up to the previous conversation. Keep the previous context in mind.\n"

    prompt += f"\nGiven the following document excerpts:\n\n{chr(10).join(context)}\n\n"
    prompt += (
        f"User: {question}\n\n"
        "Respond in the following JSON format:\n"
        "```json\n"
        "{\n"
        "  \"answer\": \"Your full, formatted answer with bullets or numbered steps.\",\n"
        "  \"follow_up_required\": true/false,\n"
        "  \"relevant_document_found\": true/false\n"
        "}\n"
        "```\n"
    )

    print("Prompt preview:", prompt[:500], flush=True)

    # Defaults
    answer = "No relevant answer found"
    follow_up_required = False
    relevant_document_found = False
    input_tokens = 0
    output_tokens = 0
    response_time = 0

    try:
        start_time_llm = time.time()
        response_obj = model.generate_content(prompt)
        response_time = time.time() - start_time_llm

        if response_obj and hasattr(response_obj, 'text'):
            raw_response = response_obj.text.strip()

            match = re.search(r'```json\s*(\{.*?\})\s*```', raw_response, re.DOTALL)
            if match:
                response_data = json.loads(match.group(1))
                answer = response_data.get("answer", "No relevant answer found")
                follow_up_required = response_data.get("follow_up_required", False)
                relevant_document_found = response_data.get("relevant_document_found", False)
            else:
                answer = raw_response

        if response_obj and hasattr(response_obj, 'usage_metadata'):
            input_tokens = getattr(response_obj.usage_metadata, 'prompt_token_count', 0)
            output_tokens = getattr(response_obj.usage_metadata, 'candidates_token_count', 0)
    except Exception as e:
        answer = f"An error occurred during response generation: {e}"
        print(answer, flush=True)

    print("Exiting query_agent_with_faiss_hnsw with answer:", answer[:80], "...", flush=True)
    return answer, input_tokens, output_tokens, response_time, most_relevant_document, follow_up_required, relevant_document_found










# async def query_agent(question, client_id, client_secret, tenant_id, site_id, chat_history=None):
#     global cached_documents
#     print("Entering query_agent with question:", question, flush=True)

#     if not cached_documents:
#         print("cached_documents is empty, attempting to load...", flush=True)
#         try:
#             await load_documents(client_id, client_secret, tenant_id, site_id)
#             print("load_documents completed.", flush=True)
#         except Exception as e:
#             error_message = f"Error loading documents: {e}"
#             print(error_message, flush=True)
#             return error_message, 0, 0, 0, None, False, False

#     if not cached_documents:
#         return "No documents found in the folder.", 0, 0, 0, None, False, False

#     is_follow_up = is_follow_up_question(question)
#     context = []
#     most_relevant_document = None

#     try:
#         question_embedding = get_embedding(question)
#         if not question_embedding:
#             return "Could not generate embedding for the question.", 0, 0, 0, None, False, False
#     except Exception as e:
#         return f"Error getting question embedding: {e}", 0, 0, 0, None, False, False

#     relevant_docs_with_similarity = []

#     for file_id, chunks in cached_documents.items():
#         best_chunk = None
#         highest_sim = -1
#         for chunk in chunks:
#             chunk_embedding = get_embedding(chunk)
#             if chunk_embedding is not None:
#                 similarity = np.dot(question_embedding, chunk_embedding)
#                 if similarity > highest_sim:
#                     highest_sim = similarity
#                     best_chunk = chunk
#         if best_chunk:
#             relevant_docs_with_similarity.append({
#                 "file_id": file_id,
#                 "chunk": best_chunk,
#                 "similarity": highest_sim
#             })

#     # Sort and select top K
#     top_k = 5
#     top_k_docs = sorted(relevant_docs_with_similarity, key=lambda x: x['similarity'], reverse=True)[:top_k]

#     if top_k_docs:
#         most_relevant_document = top_k_docs[0]["file_id"]
#         for doc in top_k_docs:
#             context.append(f"Document: {doc['file_id']}\n{doc['chunk']}")
#     else:
#         context.append("No relevant documents found.")

#     # Prompt setup
#     prompt = (
#         "You are an intelligent assistant helping IT support agents troubleshoot user issues using internal technical documentation.\n"
#         "Always assume the user is an IT support team member, not an end-user.\n\n"
#         "Given the query and the retrieved context:\n"
#         "- If the query is unclear or too broad, ask one clarifying question.\n"
#         "- Present 2–3 relevant options to guide the user (e.g., possible causes, tools to check, affected systems).\n"
#         "- If the query is clear, provide step-by-step troubleshooting steps or relevant document excerpts.\n"
#         "- Do not say 'contact IT support'; instead, provide helpful technical direction.\n"
#         "- Keep responses concise, clear, and structured (bullets or numbered steps).\n\n"
#     )

#     if chat_history:
#         for entry in chat_history[-2:]:
#             prompt += f"User: {entry['question']}\nBot: {entry['answer']}\n---\n"

#     if is_follow_up:
#         prompt += "\nThis is a follow-up to the previous conversation. Keep the previous context in mind.\n"

#     prompt += f"\nGiven the following document excerpts:\n\n{'\n\n'.join(context)}\n\n"
#     prompt += (
#         f"User: {question}\n\n"
#         "Respond in the following JSON format:\n"
#         "```json\n"
#         "{\n"
#         "  \"answer\": \"Your full, formatted answer with bullets or numbered steps.\",\n"
#         "  \"follow_up_required\": true/false,\n"
#         "  \"relevant_document_found\": true/false\n"
#         "}\n"
#         "```\n"
#     )

#     print("Prompt preview:", prompt[:500], flush=True)

#     # Defaults
#     answer = "No relevant answer found"
#     follow_up_required = False
#     relevant_document_found = False
#     input_tokens = 0
#     output_tokens = 0
#     response_time = 0

#     try:
#         start_time_llm = time.time()
#         response_obj = model.generate_content(prompt)
#         response_time = time.time() - start_time_llm

#         if response_obj and hasattr(response_obj, 'text'):
#             raw_response = response_obj.text.strip()

#             match = re.search(r'```json\s*(\{.*?\})\s*```', raw_response, re.DOTALL)
#             if match:
#                 response_data = json.loads(match.group(1))
#                 answer = response_data.get("answer", "No relevant answer found")
#                 follow_up_required = response_data.get("follow_up_required", False)
#                 relevant_document_found = response_data.get("relevant_document_found", False)
#             else:
#                 answer = raw_response

#         if response_obj and hasattr(response_obj, 'usage_metadata'):
#             input_tokens = getattr(response_obj.usage_metadata, 'prompt_token_count', 0)
#             output_tokens = getattr(response_obj.usage_metadata, 'candidates_token_count', 0)
#     except Exception as e:
#         answer = f"An error occurred during response generation: {e}"
#         print(answer, flush=True)

#     print("Exiting query_agent with answer:", answer[:80], "...", flush=True)
#     return answer, input_tokens, output_tokens, response_time, most_relevant_document, follow_up_required, relevant_document_found









 